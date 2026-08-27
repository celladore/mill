import asyncio
import io
from datetime import UTC, datetime, timedelta
from types import SimpleNamespace

import pytest
from docx import Document
from fastapi import HTTPException
from services import text_service
from services.retention_service import RetentionService
from services.artifact_storage_service import ArtifactMetadata


class TextCollection:
    def __init__(self):
        self.records = []
        self.last_query = None

    async def insert_one(self, record):
        self.records.append(record)

    async def find_one(self, query):
        self.last_query = query
        return next(
            (
                record
                for record in self.records
                if all(record.get(key) == value for key, value in query.items())
            ),
            None,
        )


class AsyncCursor:
    def __init__(self, records):
        self.records = iter(records)

    def __aiter__(self):
        return self

    async def __anext__(self):
        try:
            return next(self.records)
        except StopIteration:
            raise StopAsyncIteration


class ExpiringCollection:
    def __init__(self, records):
        self.records = records
        self.query = None
        self.updated = []

    def find(self, query):
        self.query = query
        return AsyncCursor(self.records)

    async def update_one(self, query, update):
        self.updated.append((query, update))
        return SimpleNamespace(modified_count=1)


def _mock_artifacts(monkeypatch):
    uploaded = {}

    async def upload(path, *, conversion_id, kind, user_id, content_type):
        uploaded[conversion_id] = path.read_bytes()
        now = datetime.now(UTC)
        return ArtifactMetadata(
            f"{kind}/{conversion_id}",
            content_type,
            len(uploaded[conversion_id]),
            "a" * 64,
            now,
            now + timedelta(days=7),
        )

    async def exists(_blob_name):
        return True

    monkeypatch.setattr(text_service.ArtifactStorageService, "upload", upload)
    monkeypatch.setattr(
        "services.artifact_record_service.ArtifactStorageService.exists", exists
    )
    return uploaded


def _docx_bytes(text):
    document = Document()
    document.add_heading("Matrix", level=1)
    document.add_paragraph(text)
    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


@pytest.mark.parametrize("source_format", ["md", "html", "txt", "docx"])
@pytest.mark.parametrize("target_format", ["md", "html", "txt", "docx"])
def test_complete_text_capability_matrix(source_format, target_format):
    sources = {
        "md": b"# Matrix\n\nDeterministic words",
        "html": b"<h1>Matrix</h1><p>Deterministic words</p>",
        "txt": b"Matrix\n\nDeterministic words",
        "docx": _docx_bytes("Deterministic words"),
    }
    canonical = text_service._to_html(sources[source_format], source_format)
    output = text_service._render_output(canonical, target_format)
    assert output
    if target_format == "docx":
        document = Document(io.BytesIO(output))
        rendered = " ".join(paragraph.text for paragraph in document.paragraphs)
    else:
        rendered = output.decode("utf-8")
    assert "Matrix" in rendered
    assert "Deterministic words" in rendered


def test_markdown_to_docx_is_deterministic_owner_scoped_and_downloadable(
    monkeypatch, tmp_path
):
    collection = TextCollection()
    database = SimpleNamespace(text_conversions=collection)
    monkeypatch.setattr(text_service, "TEMP_DIR", tmp_path)
    monkeypatch.setattr(text_service.Database, "get_db", lambda: database)
    uploaded = _mock_artifacts(monkeypatch)

    result = asyncio.run(
        text_service.TextService.process_text_file(
            b"# Release notes\n\n- Added private history\n- Added DOCX output",
            "release.md",
            "docx",
            "owner-1",
        )
    )

    assert result.success is True
    assert result.original_format == "md"
    assert result.target_format == "docx"
    blob_name, media_type, record = asyncio.run(
        text_service.TextService.get_output(result.id, "owner-1")
    )
    assert collection.last_query == {"id": result.id, "user_id": "owner-1"}
    assert blob_name == f"text/{result.id}"
    assert "wordprocessingml" in media_type
    document = Document(io.BytesIO(uploaded[result.id]))
    assert "Release notes" in " ".join(
        paragraph.text for paragraph in document.paragraphs
    )
    assert record["target_format"] == "docx"
    assert record["user_id"] == "owner-1"
    assert "output_path" not in result.model_dump()
    assert "expires_at" not in result.model_dump()


def test_html_is_sanitized_before_markdown_export(monkeypatch, tmp_path):
    collection = TextCollection()
    monkeypatch.setattr(text_service, "TEMP_DIR", tmp_path)
    monkeypatch.setattr(
        text_service.Database,
        "get_db",
        lambda: SimpleNamespace(text_conversions=collection),
    )
    uploaded = _mock_artifacts(monkeypatch)

    result = asyncio.run(
        text_service.TextService.process_text_file(
            (
                b"<h1>Safe</h1><script>alert('no')</script>"
                b"<a href='javascript:bad'>link</a>"
            ),
            "source.html",
            "md",
            "owner-1",
        )
    )
    output = uploaded[result.id].decode("utf-8")
    assert "# Safe" in output
    assert "<script" not in output
    assert "javascript:" not in output


def test_cancellation_after_render_removes_temporary_output(monkeypatch, tmp_path):
    class CancellingCollection(TextCollection):
        async def insert_one(self, _record):
            raise asyncio.CancelledError

    collection = CancellingCollection()
    monkeypatch.setattr(text_service, "TEMP_DIR", tmp_path)
    monkeypatch.setattr(
        text_service.Database,
        "get_db",
        lambda: SimpleNamespace(text_conversions=collection),
    )
    _mock_artifacts(monkeypatch)

    async def delete_best_effort(_blob_name, _context):
        return True

    monkeypatch.setattr(
        text_service.ArtifactStorageService,
        "delete_best_effort",
        delete_best_effort,
    )

    with pytest.raises(asyncio.CancelledError):
        asyncio.run(
            text_service.TextService.process_text_file(
                b"# Cancel me", "cancel.md", "html", "owner-1"
            )
        )

    assert list(tmp_path.iterdir()) == []


def test_rejects_unsupported_and_non_utf8_inputs(monkeypatch, tmp_path):
    monkeypatch.setattr(text_service, "TEMP_DIR", tmp_path)

    with pytest.raises(HTTPException) as unsupported:
        asyncio.run(
            text_service.TextService.process_text_file(
                b"content", "notes.rtf", "txt", "owner-1"
            )
        )
    assert unsupported.value.status_code == 400

    with pytest.raises(HTTPException) as encoding:
        asyncio.run(
            text_service.TextService.process_text_file(
                b"\xff\xfe", "notes.txt", "md", "owner-1"
            )
        )
    assert encoding.value.status_code == 400


def test_text_retention_removes_blob_but_retains_history(monkeypatch):
    collection = ExpiringCollection(
        [
            {
                "id": "expired-1",
                "artifact_blob_name": "text/expired-1",
                "artifact_available": True,
                "artifact_expires_at": datetime.now(UTC) - timedelta(seconds=1),
            }
        ]
    )
    deleted = []

    async def delete(blob_name):
        deleted.append(blob_name)

    monkeypatch.setattr(
        "services.retention_service.ArtifactStorageService.delete", delete
    )
    monkeypatch.setattr(
        "services.retention_service.Database.get_db",
        lambda: SimpleNamespace(text_conversions=collection),
    )

    expired = asyncio.run(RetentionService.expire_text_conversions())

    assert expired == 1
    assert deleted == ["text/expired-1"]
    assert collection.updated[0][0] == {"id": "expired-1", "artifact_available": True}
    assert collection.updated[0][1]["$set"]["artifact_available"] is False
    assert "$lte" in collection.query["artifact_expires_at"]
