"""Deterministic Markdown, HTML, plain-text, and DOCX transformations."""

import asyncio
import html
import io
import logging
import uuid
import zipfile
from datetime import UTC, datetime, timedelta
from pathlib import Path

import bleach
import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from config import CONVERSION_RETENTION_SECONDS, MAX_FILE_SIZE, TEMP_DIR
from database import Database
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from fastapi import HTTPException
from markdownify import markdownify
from models import TextConversionResult

from utils.security import sanitize_filename, validate_file_path, validate_target_format
from services.artifact_storage_service import ArtifactStorageService
from services.artifact_record_service import ArtifactRecordService

logger = logging.getLogger(__name__)

TEXT_FORMATS = {"md": "md", "html": "html", "txt": "txt", "docx": "docx"}
TEXT_EXTENSIONS = {
    ".md": "md",
    ".markdown": "md",
    ".html": "html",
    ".htm": "html",
    ".txt": "txt",
    ".docx": "docx",
}
MEDIA_TYPES = {
    "md": "text/markdown; charset=utf-8",
    "html": "text/html; charset=utf-8",
    "txt": "text/plain; charset=utf-8",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
ALLOWED_TAGS = {
    "a",
    "blockquote",
    "br",
    "code",
    "del",
    "em",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "hr",
    "li",
    "ol",
    "p",
    "pre",
    "strong",
    "table",
    "tbody",
    "td",
    "th",
    "thead",
    "tr",
    "ul",
}
ALLOWED_ATTRIBUTES = {"a": ["href", "title"]}
MAX_DOCX_UNCOMPRESSED_BYTES = MAX_FILE_SIZE * 8


def _safe_html(value: str) -> str:
    soup = BeautifulSoup(value, "html.parser")
    for unsafe in soup.find_all(["script", "style"]):
        unsafe.decompose()
    return bleach.clean(
        str(soup),
        tags=ALLOWED_TAGS,
        attributes=ALLOWED_ATTRIBUTES,
        protocols={"http", "https", "mailto"},
        strip=True,
    )


def _decode_utf8(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise HTTPException(
            status_code=400, detail="Text files must use UTF-8 encoding"
        ) from exc


def _docx_to_html(content: bytes) -> str:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            total_size = sum(item.file_size for item in archive.infolist())
            if total_size > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise HTTPException(
                    status_code=413,
                    detail="DOCX expands beyond the safe processing limit",
                )
        document = Document(io.BytesIO(content))
    except HTTPException:
        raise
    except (ValueError, zipfile.BadZipFile, PackageNotFoundError) as exc:
        raise HTTPException(
            status_code=400, detail="The DOCX file is invalid or damaged"
        ) from exc

    blocks = []
    for paragraph in document.paragraphs:
        text = html.escape(paragraph.text)
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style else "").lower()
        if style.startswith("heading"):
            level = next((n for n in range(1, 7) if str(n) in style), 2)
            blocks.append(f"<h{level}>{text}</h{level}>")
        elif "list" in style:
            blocks.append(f"<ul><li>{text}</li></ul>")
        else:
            blocks.append(f"<p>{text}</p>")
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = "".join(f"<td>{html.escape(cell.text)}</td>" for cell in row.cells)
            rows.append(f"<tr>{cells}</tr>")
        blocks.append(f"<table><tbody>{''.join(rows)}</tbody></table>")
    return _safe_html("\n".join(blocks))


def _to_html(content: bytes, source_format: str) -> str:
    if source_format == "docx":
        return _docx_to_html(content)
    text = _decode_utf8(content)
    if source_format == "md":
        return _safe_html(markdown.markdown(text, extensions=["extra", "sane_lists"]))
    if source_format == "txt":
        paragraphs = [
            f"<p>{html.escape(part)}</p>"
            for part in text.replace("\r\n", "\n").split("\n\n")
            if part
        ]
        return "\n".join(paragraphs)
    return _safe_html(text)


def _append_inline(paragraph, node) -> None:
    if isinstance(node, NavigableString):
        paragraph.add_run(str(node))
        return
    if not isinstance(node, Tag):
        return
    if node.name == "br":
        paragraph.add_run().add_break()
        return
    run = paragraph.add_run(node.get_text())
    run.bold = node.name in {"strong", "b"}
    run.italic = node.name in {"em", "i"}


def _html_to_docx(value: str) -> bytes:
    document = Document()
    soup = BeautifulSoup(value, "html.parser")
    for node in soup.find_all(
        ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "blockquote", "table"],
        recursive=True,
    ):
        if node.find_parent(["li", "table"]) and node.name != "table":
            continue
        if node.name == "table":
            rows = node.find_all("tr")
            if not rows:
                continue
            width = max(len(row.find_all(["td", "th"])) for row in rows)
            table = document.add_table(rows=len(rows), cols=width)
            for row_index, row in enumerate(rows):
                for col_index, cell in enumerate(row.find_all(["td", "th"])):
                    table.cell(row_index, col_index).text = cell.get_text(
                        " ", strip=True
                    )
            continue
        style = None
        if node.name.startswith("h"):
            style = f"Heading {node.name[1]}"
        elif node.name == "li":
            style = "List Bullet"
        elif node.name == "pre":
            style = "No Spacing"
        paragraph = document.add_paragraph(style=style)
        for child in node.children:
            _append_inline(paragraph, child)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _render_output(canonical_html: str, target_format: str) -> bytes:
    if target_format == "html":
        return ('<!doctype html>\n<meta charset="utf-8">\n' + canonical_html).encode(
            "utf-8"
        )
    if target_format == "md":
        return markdownify(canonical_html, heading_style="ATX").strip().encode("utf-8")
    if target_format == "txt":
        text = BeautifulSoup(canonical_html, "html.parser").get_text("\n", strip=True)
        return (text + "\n").encode("utf-8")
    return _html_to_docx(canonical_html)


class TextService:
    @staticmethod
    async def process_text_file(
        file_content: bytes, filename: str, target_format: str, user_id: str
    ) -> TextConversionResult:
        if len(file_content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File size exceeds {MAX_FILE_SIZE // (1024 * 1024)}MB limit",
            )
        safe_filename = sanitize_filename(filename)
        try:
            original_format = TEXT_EXTENSIONS[Path(safe_filename).suffix.lower()]
            target_format = validate_target_format(target_format, TEXT_FORMATS)
        except KeyError as exc:
            raise HTTPException(
                status_code=400,
                detail="Supported text inputs are Markdown, HTML, plain text, and DOCX",
            ) from exc
        except ValueError as exc:
            raise HTTPException(status_code=400, detail=str(exc)) from exc

        conversion_id = str(uuid.uuid4())
        output_path = TEMP_DIR / f"{conversion_id}.{target_format}"
        validate_file_path(TEMP_DIR, output_path)
        expires_at = datetime.now(UTC) + timedelta(seconds=CONVERSION_RETENTION_SECONDS)

        try:

            def _convert_and_write():
                rendered = _render_output(
                    _to_html(file_content, original_format), target_format
                )
                output_path.write_bytes(rendered)
                return rendered

            conversion_task = asyncio.create_task(asyncio.to_thread(_convert_and_write))
            try:
                output = await asyncio.shield(conversion_task)
            except asyncio.CancelledError:
                try:
                    await conversion_task
                except Exception:
                    pass
                finally:
                    output_path.unlink(missing_ok=True)
                raise
            result = TextConversionResult(
                id=conversion_id,
                filename=Path(safe_filename).stem,
                original_format=original_format,
                target_format=target_format,
                success=True,
                file_size_kb=round(len(output) / 1024, 2),
            )
            artifact = None
            try:
                artifact = await ArtifactStorageService.upload(
                    output_path,
                    conversion_id=conversion_id,
                    kind="text",
                    user_id=user_id,
                    content_type=MEDIA_TYPES[target_format],
                )
                persisted = result.model_dump()
                persisted.update(
                    user_id=user_id,
                    expires_at=expires_at,
                )
                persisted.update(artifact.as_record())
                await Database.get_db().text_conversions.insert_one(persisted)
            except Exception:
                if artifact:
                    await ArtifactStorageService.delete_best_effort(
                        artifact.blob_name, f"text conversion {conversion_id}"
                    )
                output_path.unlink(missing_ok=True)
                raise
            output_path.unlink(missing_ok=True)
            return result
        except HTTPException:
            raise
        except Exception as exc:
            output_path.unlink(missing_ok=True)
            logger.error("Text conversion %s failed", conversion_id, exc_info=True)
            raise HTTPException(
                status_code=500, detail="The text file could not be converted"
            ) from exc

    @staticmethod
    async def get_output(conversion_id: str, user_id: str):
        record = await ArtifactRecordService.get_download(
            Database.get_db().text_conversions, conversion_id, user_id
        )
        return (
            record["artifact_blob_name"],
            MEDIA_TYPES[record["target_format"]],
            record,
        )
